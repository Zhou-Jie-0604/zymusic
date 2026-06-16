#!/usr/bin/env python3
"""
Modify zy音乐（1）.docx:
1. Replace "（这里需要插入XX图）" placeholders with generated images
2. Convert inline text formulas to Word OMML equation format
"""

import os
import re
import io
from copy import deepcopy
from io import BytesIO

import matplotlib
matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

# ============================================================
# Constants
# ============================================================
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

NSMAP = {
    'm': MATH_NS,
    'w': WORD_NS,
}

# Register namespace for clean XML output
etree.register_namespace('m', MATH_NS)

# ============================================================
# OMML Helper Functions
# ============================================================

def _m(tag):
    """Create an element in the math namespace."""
    return etree.Element(f'{{{MATH_NS}}}{tag}', nsmap=NSMAP)

def _m_sub(tag, parent):
    """Create a sub-element in the math namespace."""
    return etree.SubElement(parent, f'{{{MATH_NS}}}{tag}', nsmap=NSMAP)

def make_math_run(text):
    """Create an <m:r> element with <m:t> inside."""
    mr = _m('r')
    mt = _m_sub('t', mr)
    mt.text = text
    return mr

def make_omath(*children):
    """Create an <m:oMath> wrapper."""
    om = _m('oMath')
    for child in children:
        om.append(child)
    return om

def make_omath_para(*children):
    """Create an <m:oMathPara> wrapper for display math."""
    op = _m('oMathPara')
    for child in children:
        op.append(child)
    return op

def make_delim(beg_char, end_char, *content):
    """Create parenthesized expression <m:d> with beg/end chars."""
    md = _m('d')
    # Delimiter properties
    dPr = _m_sub('dPr', md)
    beg = _m_sub('begChr', dPr)
    beg.set(f'{{{MATH_NS}}}val', beg_char)
    end = _m_sub('endChr', dPr)
    end.set(f'{{{MATH_NS}}}val', end_char)
    # Content
    me = _m_sub('e', md)
    for child in content:
        me.append(child)
    return md

def make_paren(*content):
    """Create parenthesized expression with standard parentheses."""
    return make_delim('(', ')', *content)

def make_floor(*content):
    """Create floor-bracketed expression ⌊ ⌋."""
    return make_delim('⌊', '⌋', *content)

def make_frac(num_children, den_children):
    """Create a fraction <m:f>."""
    mf = _m('f')
    # Numerator
    num = _m_sub('num', mf)
    for child in num_children:
        num.append(child)
    # Denominator
    den = _m_sub('den', mf)
    for child in den_children:
        den.append(child)
    return mf

def make_sub(base_children, sub_children):
    """Create subscript <m:sSub>."""
    ss = _m('sSub')
    e = _m_sub('e', ss)
    for child in base_children:
        e.append(child)
    s = _m_sub('sub', ss)
    for child in sub_children:
        s.append(child)
    return ss

def make_simple_formula(text):
    """Create a simple formula from plain text (no special structure)."""
    return make_omath(make_math_run(text))

def make_formula_with_parts(*parts):
    """
    Create formula from parts where each part is either:
    - a string (becomes a math run)
    - an etree Element (used directly)
    """
    children = []
    for part in parts:
        if isinstance(part, str):
            children.append(make_math_run(part))
        else:
            children.append(part)
    return make_omath(*children)

# ============================================================
# Specific Formula Builders
# ============================================================

def build_sigma_selection(music_id=True):
    """Build σ(musicId=m.id ∧ parentId=0)(C) or σ(parentId=p.id)(C)."""
    parts = []
    parts.append(make_math_run('σ'))
    if music_id:
        parts.append(make_paren(
            make_math_run('musicId=m.id ∧ parentId=0')
        ))
    else:
        parts.append(make_paren(
            make_math_run('parentId=p.id')
        ))
    parts.append(make_paren(make_math_run('C')))
    return make_omath(*parts)

def build_element_of():
    """Build c ∈ C."""
    return make_omath(
        make_math_run('c ∈ C')
    )

def build_big_o(expr):
    """Build O(expr) notation."""
    return make_omath(
        make_math_run('O'),
        make_paren(make_math_run(expr))
    )

def build_set_notation(items):
    """Build {item1, item2, ...} set notation."""
    return make_omath(
        make_delim('{', '}', make_math_run(items))
    )

def build_state_flip():
    """Build f(已赞)=未赞."""
    return make_omath(
        make_math_run('f'),
        make_paren(make_math_run('已赞')),
        make_math_run('=未赞')
    )

def build_state_flip2():
    """Build f(未赞)=已赞."""
    return make_omath(
        make_math_run('f'),
        make_paren(make_math_run('未赞')),
        make_math_run('=已赞')
    )

def build_interval():
    """Build 0 ≤ start ≤ end < L."""
    return make_omath(
        make_math_run('0 ≤ start ≤ end < L')
    )

def build_range_formula():
    """Build (end − start + 1)."""
    return make_omath(
        make_paren(
            make_math_run('end − start + 1')
        )
    )

def build_floor_formula():
    """Build start = ⌊(t / T) × L⌋."""
    return make_omath(
        make_math_run('start = '),
        make_floor(
            make_frac(
                [make_math_run('t')],
                [make_math_run('T')]
            ),
            make_math_run(' × L')
        )
    )

def build_mode_formula():
    """Build mode_next = (mode_current + 1) mod 3."""
    return make_omath(
        make_math_run('mode'),
        make_sub([make_math_run('next')], [make_math_run('')]),
        make_math_run('= (mode'),
        make_sub([make_math_run('current')], [make_math_run('')]),
        make_math_run(' + 1) mod 3')
    )

def build_mode_formula_simple():
    """Build mode_next = (mode_current + 1) mod 3 (simplified)."""
    return make_omath(
        make_math_run('mode_next = (mode_current + 1) mod 3')
    )

def build_interval_0_n1():
    """Build [0, n-1]."""
    return make_omath(
        make_delim('[', ']', make_math_run('0, n−1'))
    )

def build_fraction_simple(num, den):
    """Build simple fraction like 1/n."""
    return make_omath(
        make_frac(
            [make_math_run(num)],
            [make_math_run(den)]
        )
    )

def build_expected_samples():
    """Build n/(n-1)."""
    return make_omath(
        make_frac(
            [make_math_run('n')],
            [make_math_run('n−1')]
        )
    )

def build_base_scale():
    """Build baseScale = max(size/iw, size/ih)."""
    return make_omath(
        make_math_run('baseScale = max(size/iw, size/ih)')
    )

def build_image_dims():
    """Build iw×ih."""
    return make_omath(
        make_math_run('iw × ih')
    )

def build_scale_formula():
    """Build scale = baseScale × zoomScale."""
    return make_omath(
        make_math_run('scale = baseScale × zoomScale')
    )

def build_anchor_x():
    """Build anchorX = (size − iw×scale)/2 + dx."""
    return make_omath(
        make_math_run('anchorX = (size − iw × scale)/2 + dx')
    )

def build_anchor_y():
    """Build anchorY = (size − ih×scale)/2 + dy."""
    return make_omath(
        make_math_run('anchorY = (size − ih × scale)/2 + dy')
    )

def build_centering_formula():
    """Build (size − iw×scale)/2."""
    return make_omath(
        make_paren(
            make_math_run('size − iw × scale'),
        ),
        make_math_run('/2')
    )

def build_center_point():
    """Build (size/2, size/2)."""
    return make_omath(
        make_paren(
            make_math_run('size/2, size/2')
        )
    )

# ============================================================
# Paragraph Formula Replacement Logic
# ============================================================

class FormulaReplacement:
    """Represents a formula to insert within a paragraph."""
    def __init__(self, text_before, omath_element):
        self.text_before = text_before  # text that comes before the formula
        self.omath = omath_element  # the OMML element

def replace_formulas_in_paragraph(para, replacements):
    """
    Rebuild a paragraph with text runs and OMML formula elements.

    replacements: list of tuples (plain_text_or_FormulaReplacement, ...)
    Each item is either a string (plain text) or a FormulaReplacement.
    FormulaReplacement.text_before is prepended as text, then the OMML is inserted.

    For simpler use: pass a list where each element is either:
    - a str: added as a text run
    - an Element (OMLL): inserted as an equation
    """
    # Clear existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    for item in replacements:
        if isinstance(item, str):
            # Add text run
            run = para.add_run(item)
            run.font.size = Pt(12)
        elif isinstance(item, FormulaReplacement):
            # Add text before formula
            if item.text_before:
                run = para.add_run(item.text_before)
                run.font.size = Pt(12)
            # Insert OMML equation
            para._element.append(item.omath)
        else:
            # Assume it's an OMML element
            para._element.append(item)


def clear_para(para):
    """Remove all runs from a paragraph."""
    for run in para.runs:
        run._element.getparent().remove(run._element)


# ============================================================
# Image Generation Functions
# ============================================================

def generate_layer_architecture():
    """Generate the 5-layer system architecture diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    layers = [
        ('页面表示层\n(Presentation Layer)', '#E3F2FD', 'JSP动态模板\nHTML生成'),
        ('请求分发控制层\n(Controller Layer)', '#BBDEFB', 'Servlet请求分发\n参数校验 · 会话管理'),
        ('业务规则逻辑层\n(Business Logic Layer)', '#90CAF9', '认证服务 · 评论服务\n上传服务 · 检索服务'),
        ('数据持久访问层\n(Data Access Layer)', '#64B5F6', 'JDBC原生接口\nDAO模式封装'),
        ('物理数据存储层\n(Data Storage Layer)', '#42A5F5', 'SQLite数据库\n文件系统存储'),
    ]

    colors_border = ['#1565C0', '#1976D2', '#1E88E5', '#2196F3', '#2E7D32']
    box_width = 8
    box_height = 1.3
    start_y = 8
    gap = 0.3

    for i, (title, bg_color, desc) in enumerate(layers):
        y = start_y - i * (box_height + gap)
        # Main box
        rect = FancyBboxPatch(
            (1, y - box_height), box_width, box_height,
            boxstyle="round,pad=0.1",
            facecolor=bg_color, edgecolor=colors_border[i],
            linewidth=2.5, alpha=0.95
        )
        ax.add_patch(rect)

        # Title
        ax.text(5, y - box_height/2 - 0.15, title,
                ha='center', va='center', fontsize=14, fontweight='bold',
                color='#0D47A1', fontfamily='Microsoft YaHei')
        # Subtitle
        ax.text(5, y - box_height/2 - 0.55, desc,
                ha='center', va='center', fontsize=10,
                color='#37474F', fontfamily='Microsoft YaHei', style='italic')

        # Down arrows between layers
        if i < len(layers) - 1:
            arrow_y = y - box_height - gap/2
            ax.annotate('', xy=(5, arrow_y - gap/2), xytext=(5, arrow_y + gap/2),
                       arrowprops=dict(arrowstyle='->', lw=2.5, color='#F44336'))
            ax.text(6.5, arrow_y, '仅调用下层接口', fontsize=8, color='#F44336',
                   fontfamily='Microsoft YaHei', fontweight='bold')

    # Title
    ax.text(5, 9.5, 'ZY音乐系统 — 五层分层数据处理架构',
            ha='center', va='center', fontsize=17, fontweight='bold',
            color='#1A237E', fontfamily='Microsoft YaHei')

    # Left side annotation
    ax.text(0.3, 4.5, '单向\n依赖', ha='center', va='center',
            fontsize=11, color='#D32F2F', fontweight='bold',
            fontfamily='Microsoft YaHei', rotation=90)

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


def generate_upload_flowchart():
    """Generate the music upload & format conversion flowchart."""
    fig, ax = plt.subplots(1, 1, figsize=(11, 8.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 11)
    ax.axis('off')

    # Define boxes: (x, y, width, height, text, color, border_color)
    boxes = [
        (4, 10, 3, 0.8, '客户端上传\n(多部件表单)', '#E8EAF6', '#283593'),
        (4, 8.5, 3, 0.8, '服务端接收\n提取音频/封面/元数据', '#E3F2FD', '#1565C0'),
        (4, 7.0, 3, 0.8, '多层校验\n(容量≤100MB · 格式白名单\n· MIME类型检测)', '#FFF3E0', '#E65100'),

        # Decision diamond
        (4.5, 5.8, 2, 0.8, '校验\n通过?', '#FFEBEE', '#B71C1C'),

        (0.5, 4.3, 2.8, 0.8, 'UUID生成唯一文件名\n数据库存相对路径', '#E8F5E9', '#2E7D32'),
        (4.5, 4.3, 3.2, 0.8, '格式检测\n是否为FLAC?', '#F3E5F5', '#6A1B9A'),

        # FLAC branch
        (0.3, 2.5, 3.2, 0.8, 'FFmpeg转换\n→320kbps MP3', '#E0F7FA', '#00695C'),
        (0.3, 1.0, 3.2, 0.8, '降级:自实现解码器\n→WAV格式缓存', '#E0F7FA', '#00695C'),
        (0.3, -0.3, 3.2, 0.8, '最终降级:原始FLAC\n直接提供播放', '#FFCDD2', '#C62828'),

        # Non-FLAC branch
        (7.5, 4.3, 3, 0.8, 'MP3/WAV等格式\n直接入库', '#C8E6C9', '#388E3C'),

        # Merge point
        (3.5, -1.2, 4, 0.8, '数据库写入完成\n客户端异步返回结果', '#ECEFF1', '#37474F'),

        # Error path
        (8, 5.8, 2.5, 0.6, '拒绝上传\n返回错误信息', '#FFCDD2', '#C62828'),
    ]

    for (x, y, w, h, text, bg, border) in boxes:
        if '校验' in text and '通过?' in text:
            # Diamond shape using a circle-like box
            rect = FancyBboxPatch(
                (x, y - h), w, h,
                boxstyle="round,pad=0.2",
                facecolor=bg, edgecolor=border, linewidth=2, alpha=0.9
            )
        else:
            rect = FancyBboxPatch(
                (x, y - h), w, h,
                boxstyle="round,pad=0.15",
                facecolor=bg, edgecolor=border, linewidth=2.2, alpha=0.9
            )
        ax.add_patch(rect)
        ax.text(x + w/2, y - h/2, text, ha='center', va='center',
               fontsize=9, fontweight='bold', color='#212121',
               fontfamily='Microsoft YaHei')

    # Arrows
    arrows = [
        # Main flow
        (5.5, 8.5, 5.5, 9.2, '#37474F'),
        (5.5, 7.0, 5.5, 7.7, '#37474F'),
        (5.5, 5.8, 5.5, 6.2, '#37474F'),
        # Yes branch
        (5.5, 5.0, 5.5, 5.0, '#2E7D32'),  # to UUID
        # To format detection
        (1.9, 4.3, 5.0, 5.0, '#37474F'),
        # Format detection - FLAC path
        (6.1, 4.3, 1.9, 3.3, '#00695C'),
        # Format detection - non-FLAC path
        (6.1, 4.3, 9.0, 5.0, '#388E3C'),
        # Non-FLAC to merge
        (9.0, 4.3, 5.5, -0.4, '#388E3C'),
        # FLAC chain
        (1.9, 2.5, 1.9, 3.5, '#00695C'),
        (1.9, 1.0, 1.9, 1.7, '#00695C'),
        (1.9, -0.3, 1.9, 0.2, '#C62828'),
        # FLAC final to merge
        (1.9, -0.3, 5.5, -0.4, '#37474F'),
        # No (reject) path
        (5.5, 5.8, 9.25, 5.8, '#C62828'),
    ]

    for (x1, y1, x2, y2, color) in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', lw=2, color=color,
                                  connectionstyle='arc3,rad=0'))

    # Labels on key arrows
    ax.text(5.8, 5.4, '是', fontsize=10, color='#2E7D32', fontweight='bold')
    ax.text(8.0, 5.4, '否', fontsize=10, color='#C62828', fontweight='bold')
    ax.text(2.2, 3.9, '是(FLAC)', fontsize=9, color='#00695C', fontweight='bold')
    ax.text(7.2, 4.3, '否(其他格式)', fontsize=9, color='#388E3C', fontweight='bold')

    # Title
    ax.text(5.5, 10.8, 'ZY音乐系统 — 音乐上传与格式转换处理流程',
            ha='center', va='center', fontsize=17, fontweight='bold',
            color='#1A237E', fontfamily='Microsoft YaHei')

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


def generate_http_range_diagram():
    """Generate HTTP Range request/response sequence diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis('off')

    # Client and Server boxes
    client_x = 2
    server_x = 9

    # Draw lifelines
    ax.plot([client_x, client_x], [0.5, 8.5], 'k--', lw=1.5, alpha=0.5)
    ax.plot([server_x, server_x], [0.5, 8.5], 'k--', lw=1.5, alpha=0.5)

    # Client box
    client_box = FancyBboxPatch(
        (client_x - 1.2, 8.2), 2.4, 0.7,
        boxstyle="round,pad=0.1", facecolor='#E3F2FD',
        edgecolor='#1565C0', linewidth=2.5
    )
    ax.add_patch(client_box)
    ax.text(client_x, 8.55, '客户端 (Browser)', ha='center', va='center',
           fontsize=13, fontweight='bold', color='#0D47A1', fontfamily='Microsoft YaHei')

    # Server box
    server_box = FancyBboxPatch(
        (server_x - 1.2, 8.2), 2.4, 0.7,
        boxstyle="round,pad=0.1", facecolor='#E8F5E9',
        edgecolor='#2E7D32', linewidth=2.5
    )
    ax.add_patch(server_box)
    ax.text(server_x, 8.55, '服务端 (Jetty)', ha='center', va='center',
           fontsize=13, fontweight='bold', color='#1B5E20', fontfamily='Microsoft YaHei')

    # ===== Interactions =====

    # 1. First request: GET without Range
    arrow_y1 = 7.2
    ax.annotate('', xy=(server_x, arrow_y1), xytext=(client_x, arrow_y1),
               arrowprops=dict(arrowstyle='->', lw=2.2, color='#37474F'))
    ax.text(5.5, arrow_y1 + 0.2, '① GET /files/music/xxx.mp3 (无Range头)',
           ha='center', fontsize=10, fontweight='bold', color='#37474F',
           fontfamily='Microsoft YaHei')

    # Response 1: 200 OK with full metadata
    arrow_y1b = 6.5
    ax.annotate('', xy=(client_x, arrow_y1b), xytext=(server_x, arrow_y1b),
               arrowprops=dict(arrowstyle='->', lw=1.5, color='#1565C0',
                              linestyle='dashed'))
    ax.text(5.5, arrow_y1b - 0.35, '← 200 OK + Content-Length: L + Content-Type',
           ha='center', fontsize=9, color='#1565C0', fontfamily='Microsoft YaHei')

    # 2. Range request for seeking
    arrow_y2 = 5.5
    ax.annotate('', xy=(server_x, arrow_y2), xytext=(client_x, arrow_y2),
               arrowprops=dict(arrowstyle='->', lw=2.2, color='#E65100'))
    ax.text(5.5, arrow_y2 + 0.23, '② GET ... + Range: bytes=start-end',
           ha='center', fontsize=10, fontweight='bold', color='#E65100',
           fontfamily='Microsoft YaHei')

    # Server processing indicator
    proc_box = FancyBboxPatch(
        (server_x - 0.3, 4.9), 0.6, 0.5,
        boxstyle="round,pad=0.05", facecolor='#FFF9C4',
        edgecolor='#F9A825', linewidth=1.5
    )
    ax.add_patch(proc_box)
    ax.text(server_x, 5.15, '解析\n校验', ha='center', va='center',
           fontsize=7, color='#E65100', fontweight='bold')

    # Response 2: 206 Partial Content
    arrow_y2b = 4.1
    ax.annotate('', xy=(client_x, arrow_y2b), xytext=(server_x, arrow_y2b),
               arrowprops=dict(arrowstyle='->', lw=1.5, color='#2E7D32',
                              linestyle='dashed'))
    ax.text(5.5, arrow_y2b - 0.35, '← 206 Partial Content + Content-Range: bytes start-end/L',
           ha='center', fontsize=9, color='#2E7D32', fontfamily='Microsoft YaHei')

    # 3. Subsequent range requests (multi-range support)
    arrow_y3 = 3.0
    ax.annotate('', xy=(server_x, arrow_y3), xytext=(client_x, arrow_y3),
               arrowprops=dict(arrowstyle='->', lw=2.2, color='#6A1B9A'))
    ax.text(5.5, arrow_y3 + 0.23, '③ Range: bytes=start1-end1, start2-end2 (多段请求)',
           ha='center', fontsize=10, fontweight='bold', color='#6A1B9A',
           fontfamily='Microsoft YaHei')

    arrow_y3b = 2.3
    ax.annotate('', xy=(client_x, arrow_y3b), xytext=(server_x, arrow_y3b),
               arrowprops=dict(arrowstyle='->', lw=1.5, color='#6A1B9A',
                              linestyle='dashed'))
    ax.text(5.5, arrow_y3b - 0.35, '← 206 + multipart/byteranges (多段数据)',
           ha='center', fontsize=9, color='#6A1B9A', fontfamily='Microsoft YaHei')

    # 4. Client-side computation box
    comp_box = FancyBboxPatch(
        (client_x - 1.5, 1.0), 3.0, 0.9,
        boxstyle="round,pad=0.1", facecolor='#FCE4EC',
        edgecolor='#C62828', linewidth=2, alpha=0.85
    )
    ax.add_patch(comp_box)
    ax.text(client_x, 1.45, '客户端计算\nstart = [(t/T) × L]\n构造Range请求头',
           ha='center', va='center', fontsize=9, fontweight='bold',
           color='#C62828', fontfamily='Microsoft YaHei')

    # Title
    ax.text(6, 9.2, 'ZY音乐系统 — HTTP范围请求与响应时序交互',
            ha='center', va='center', fontsize=17, fontweight='bold',
            color='#1A237E', fontfamily='Microsoft YaHei')

    # Legend
    legend_y = 0.3
    ax.plot([0.8, 1.5], [legend_y, legend_y], '-', lw=2.2, color='#37474F')
    ax.text(1.6, legend_y, '请求', fontsize=8, va='center', color='#37474F')
    ax.plot([3.5, 4.2], [legend_y, legend_y], '--', lw=1.5, color='#1565C0')
    ax.text(4.3, legend_y, '响应', fontsize=8, va='center', color='#1565C0')

    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


# ============================================================
# Main Document Processing
# ============================================================

def process_document(doc_path, output_path):
    """Main function: process the document."""
    print("Opening document...")
    doc = Document(doc_path)

    # ===== STEP 1: Generate and Insert Images =====
    print("\n=== STEP 1: Generating and inserting images ===")

    # Generate images
    print("Generating layer architecture diagram...")
    img1_buf = generate_layer_architecture()
    print("Generating upload flowchart...")
    img2_buf = generate_upload_flowchart()
    print("Generating HTTP Range sequence diagram...")
    img3_buf = generate_http_range_diagram()

    # Image placeholder paragraphs (0-indexed)
    # P44 -> index 44, P49 -> index 49, P64 -> index 64
    image_placeholders = {
        44: (img1_buf, '系统分层架构示意图'),
        49: (img2_buf, '音乐上传与格式转换处理流程图'),
        64: (img3_buf, 'HTTP范围请求与响应时序交互图'),
    }

    for para_idx, (img_buf, caption) in image_placeholders.items():
        para = doc.paragraphs[para_idx]
        print(f"  Processing paragraph {para_idx}: replacing '{para.text}'")

        # Clear the paragraph
        clear_para(para)

        # Add the image
        run = para.add_run()
        img_buf.seek(0)
        run.add_picture(img_buf, width=Inches(5.5))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print("Images inserted successfully.")

    # ===== STEP 2: Convert Formulas to OMML =====
    print("\n=== STEP 2: Converting formulas to OMML equations ===")

    # We'll rebuild paragraphs that contain formulas
    # Track processed paragraphs to avoid double-processing

    # Helper to get paragraph full original text
    def get_full_text(para_idx):
        return doc.paragraphs[para_idx].text

    # ---- P53: Relational algebra formulas ----
    p53 = doc.paragraphs[53]
    p53_text = get_full_text(53)
    print(f"  P53: {p53_text[:80]}...")

    # Split at sigma formulas
    # Text structure: ...节点：σ(musicId=m.id ∧ parentId=0)(C)；...集合：σ(parentId=p.id)(C)。...
    parts = re.split(r'(σ\([^)]+\)\(C\))', p53_text)
    # parts will be: [before_sigma1, sigma1, between, sigma2, after]

    clear_para(p53)
    for part in parts:
        if part.startswith('σ'):
            if 'musicId' in part:
                p53._element.append(build_sigma_selection(music_id=True))
            else:
                p53._element.append(build_sigma_selection(music_id=False))
        else:
            run = p53.add_run(part)
            run.font.size = Pt(12)

    # ---- P57: Big O notation ----
    p57 = doc.paragraphs[57]
    p57_text = get_full_text(57)
    print(f"  P57: {p57_text[:80]}...")

    # Replace O(knm)
    parts = re.split(r'(O\(knm\))', p57_text)
    clear_para(p57)
    for part in parts:
        if part == 'O(knm)':
            p57._element.append(build_big_o('knm'))
        else:
            run = p57.add_run(part)
            run.font.size = Pt(12)

    # ---- P59: State flip functions ----
    p59 = doc.paragraphs[59]
    p59_text = get_full_text(59)
    print(f"  P59: {p59_text[:80]}...")

    # Replace f(已赞)=未赞 and f(未赞)=已赞
    parts = re.split(r'(f\(已赞\)=未赞|f\(未赞\)=已赞)', p59_text)
    clear_para(p59)
    for part in parts:
        if part == 'f(已赞)=未赞':
            p59._element.append(build_state_flip())
        elif part == 'f(未赞)=已赞':
            p59._element.append(build_state_flip2())
        else:
            run = p59.add_run(part)
            run.font.size = Pt(12)

    # ---- Robust formula replacement using position-based matching ----

    def escape_regex(text):
        """Escape all regex metacharacters in a literal text string."""
        # Order matters: escape backslash first
        special_chars = r'\.^$*+?{}[]()|'
        result = ''
        for char in text:
            if char in special_chars:
                result += '\\' + char
            else:
                result += char
        return result

    def rebuild_para_with_formulas(doc, para_idx, formula_specs):
        """
        Rebuild a paragraph with formulas inserted at specified positions.

        formula_specs: list of (formula_text_string, omath_element) tuples.
        Formulas are found by literal string matching in the paragraph text.
        Positions are automatically sorted.
        """
        para = doc.paragraphs[para_idx]
        full_text = para.text

        # Find all formula positions
        found = []
        for formula_str, omath in formula_specs:
            pos = full_text.find(formula_str)
            if pos >= 0:
                found.append((pos, len(formula_str), omath, formula_str))
            else:
                print(f"    WARNING: Formula '{formula_str[:50]}...' not found in P{para_idx}")

        if not found:
            print(f"    No formulas found in P{para_idx}, skipping")
            return

        # Sort by position
        found.sort(key=lambda x: x[0])

        # Verify no overlaps
        for i in range(len(found) - 1):
            end_i = found[i][0] + found[i][1]
            start_j = found[i+1][0]
            if end_i > start_j:
                print(f"    WARNING: Formula overlap in P{para_idx}: "
                      f"'{found[i][3][:30]}' ends at {end_i}, "
                      f"'{found[i+1][3][:30]}' starts at {start_j}")

        # Build parts list
        parts_list = []
        cursor = 0
        for pos, length, omath, _ in found:
            if pos > cursor:
                parts_list.append(full_text[cursor:pos])
            parts_list.append(omath)
            cursor = pos + length

        if cursor < len(full_text):
            parts_list.append(full_text[cursor:])

        # Rebuild paragraph
        clear_para(para)
        for part in parts_list:
            if isinstance(part, str):
                if part:  # Skip empty strings
                    run = para.add_run(part)
                    run.font.size = Pt(12)
            else:
                para._element.append(part)

        print(f"  P{para_idx}: inserted {len(found)} formula(s)")

    # ---- P63: Multiple formulas ----
    print(f"  P63: {get_full_text(63)[:80]}...")
    rebuild_para_with_formulas(doc, 63, [
        ('0 ≤ start ≤ end < L', build_interval()),
        ('(end - start + 1)', build_range_formula()),
        ('start = ⌊(t / T) × L⌋', build_floor_formula()),
    ])

    # ---- P80: Playback mode formulas ----
    print(f"  P80: {get_full_text(80)[:80]}...")
    rebuild_para_with_formulas(doc, 80, [
        ('mode_next = (mode_current + 1) mod 3', build_mode_formula_simple()),
        ('[0, n-1]', build_interval_0_n1()),
        ('1/n', build_fraction_simple('1', 'n')),
        ('n/(n-1)', build_expected_samples()),
        ('O(1)', build_big_o('1')),
    ])

    # ---- P82: Image scaling formulas ----
    print(f"  P82: {get_full_text(82)[:80]}...")
    rebuild_para_with_formulas(doc, 82, [
        ('iw×ih', build_image_dims()),
        ('baseScale = max(size/iw, size/ih)', build_base_scale()),
        ('scale = baseScale × zoomScale', build_scale_formula()),
    ])

    # ---- P83: Anchor coordinate formulas (handle sub-formula overlap carefully) ----
    print(f"  P83: {get_full_text(83)[:80]}...")

    p83 = doc.paragraphs[83]
    p83_text = p83.text

    # Find the two anchor formulas and the standalone centering formula
    # anchorX = (size - iw×scale)/2 + dx  (contains the substring '(size - iw×scale)/2')
    # anchorY = (size - ih×scale)/2 + dy
    # (size - iw×scale)/2 (appears independently later)

    anchor_x_str = 'anchorX = (size - iw×scale)/2 + dx'
    anchor_y_str = 'anchorY = (size - ih×scale)/2 + dy'
    centering_str = '(size - iw×scale)/2'
    center_pt_str = '(size/2, size/2)'

    # Find positions
    ax_pos = p83_text.find(anchor_x_str)
    ay_pos = p83_text.find(anchor_y_str)

    # For the standalone centering formula, find the SECOND occurrence
    first_centering = p83_text.find(centering_str)
    second_centering = p83_text.find(centering_str, first_centering + 1)

    parts_list = []
    cursor = 0

    # Insert anchorX
    if ax_pos >= 0:
        parts_list.append(p83_text[cursor:ax_pos])
        parts_list.append(build_anchor_x())
        cursor = ax_pos + len(anchor_x_str)

    # Text between anchorX and anchorY
    if ay_pos >= 0:
        parts_list.append(p83_text[cursor:ay_pos])
        parts_list.append(build_anchor_y())
        cursor = ay_pos + len(anchor_y_str)

    # Now handle the standalone (size - iw×scale)/2
    if second_centering >= cursor:
        parts_list.append(p83_text[cursor:second_centering])
        parts_list.append(build_centering_formula())
        cursor = second_centering + len(centering_str)

    # Remaining text
    if cursor < len(p83_text):
        parts_list.append(p83_text[cursor:])

    clear_para(p83)
    for part in parts_list:
        if isinstance(part, str):
            if part:
                run = p83.add_run(part)
                run.font.size = Pt(12)
        else:
            p83._element.append(part)
    print(f"  P83: inserted formulas")

    # ---- P84: Canvas pipeline - (size/2, size/2) ----
    print(f"  P84: {get_full_text(84)[:80]}...")
    rebuild_para_with_formulas(doc, 84, [
        ('(size/2, size/2)', build_center_point()),
    ])

    # ---- P90: O(1) and O(n) in comment tree ----
    print(f"  P90: {get_full_text(90)[:80]}...")
    rebuild_para_with_formulas(doc, 90, [
        ('O(1)', build_big_o('1')),
        ('O(n)', build_big_o('n')),
    ])

    # ---- P54: O(1) in comment tree algorithm description ----
    print(f"  P54: {get_full_text(54)[:80]}...")
    rebuild_para_with_formulas(doc, 54, [
        ('O(1)', build_big_o('1')),
    ])

    # ---- P75: O(1) in notification polling complexity ----
    print(f"  P75: {get_full_text(75)[:80]}...")
    rebuild_para_with_formulas(doc, 75, [
        ('O(1)', build_big_o('1')),
    ])

    # ---- P96: n/(n-1) in conclusion ----
    print(f"  P96: {get_full_text(96)[:80]}...")
    rebuild_para_with_formulas(doc, 96, [
        ('n/(n-1)', build_expected_samples()),
    ])

    print("\n=== STEP 3: Saving modified document ===")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == '__main__':
    input_path = 'F:/ZYmusic(2)/zy音乐（1）.docx'
    output_path = 'F:/ZYmusic(2)/zy音乐（1）_modified.docx'
    process_document(input_path, output_path)
    print("\nDone!")
